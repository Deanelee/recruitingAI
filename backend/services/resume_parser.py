import io
import pdfplumber
from docx import Document
from typing import Optional


class ResumeParser:
    def parse(self, file_bytes: bytes, filename: str) -> str:
        filename_lower = filename.lower()

        if filename_lower.endswith(".pdf"):
            return self._parse_pdf(file_bytes)
        elif filename_lower.endswith(".docx"):
            return self._parse_docx(file_bytes)
        elif filename_lower.endswith(".doc"):
            return self._parse_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported file format: {filename}. Only PDF and DOCX are supported.")

    def _parse_pdf(self, file_bytes: bytes) -> str:
        text_parts = []

        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        full_text = "\n\n".join(text_parts)

        if not full_text.strip():
            raise ValueError("Could not extract text from PDF. The file may be image-based or corrupted.")

        return full_text.strip()

    def _parse_docx(self, file_bytes: bytes) -> str:
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    text_parts.append(" | ".join(row_texts))

        full_text = "\n".join(text_parts)

        if not full_text.strip():
            raise ValueError("Could not extract text from DOCX. The file may be empty or corrupted.")

        return full_text.strip()
